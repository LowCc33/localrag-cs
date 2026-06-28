#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
导出报告接口路由

提供问答报告的 PDF 和 Word 导出功能。
优先使用 python-docx 和 reportlab，失败时降级为 RTF 格式。
"""

import io
import logging
import os
from datetime import datetime
from typing import List, Optional, Dict, Any

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

# 导入会话管理器
from session_manager import session_manager

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/export", tags=["export"])


class ExportRequest(BaseModel):
    """导出请求模型"""
    session_id: str = Field(..., description="会话ID")
    title: str = Field(default="问答报告", description="报告标题")
    include_sources: bool = Field(default=True, description="是否包含来源引用")
    save_path: Optional[str] = Field(default=None, description="保存路径（服务器目录），为空则直接返回文件流")


class ExportItem(BaseModel):
    """导出问答项"""
    question: str = Field(..., description="用户问题")
    answer: str = Field(..., description="AI回答")
    sources: List[str] = Field(default=[], description="来源文档列表")


# ========== 数据准备 ==========

def _prepare_export_data(session_id: str, include_sources: bool) -> List[ExportItem]:
    """
    从会话历史中准备导出数据

    Args:
        session_id: 会话ID
        include_sources: 是否包含来源

    Returns:
        问答项列表
    """
    history = session_manager.get_history(session_id, limit=50)
    items = []
    i = 0
    while i < len(history) - 1:
        if history[i]["role"] == "user" and history[i + 1]["role"] == "assistant":
            items.append(ExportItem(
                question=history[i]["content"],
                answer=history[i + 1]["content"],
                sources=[]  # 来源信息不在历史中，留空
            ))
            i += 2
        else:
            i += 1
    return items


# ========== Word 导出 ==========

def _export_docx(items: List[ExportItem], title: str) -> bytes:
    """
    导出为 Word 文档（使用 python-docx）

    Args:
        items: 问答项列表
        title: 报告标题

    Returns:
        Word 文档字节流
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 标题
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"导出日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # 空行

    # 分隔线
    doc.add_paragraph("─" * 50)

    # 问答内容
    for idx, item in enumerate(items, 1):
        # 问题
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{idx}: {item.question}")
        q_run.bold = True
        q_run.font.size = Pt(12)

        # 答案
        a_para = doc.add_paragraph(item.answer)
        a_para.paragraph_format.space_after = Pt(6)

        # 来源
        if item.sources:
            s_para = doc.add_paragraph()
            s_run = s_para.add_run(f"来源：{'、'.join(item.sources)}")
            s_run.font.size = Pt(9)
            s_run.font.color.rgb = RGBColor(100, 100, 100)
            s_run.italic = True

        # 分隔线
        doc.add_paragraph("─" * 50)

    # 页脚
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"—— 由 LocalRAG-CS 智能客服系统生成 ——")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)

    # 输出到字节流
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ========== PDF 导出 ==========

def _export_pdf(items: List[ExportItem], title: str) -> bytes:
    """
    导出为 PDF 文档（使用 reportlab）

    Args:
        items: 问答项列表
        title: 报告标题

    Returns:
        PDF 文档字节流
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm, cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, PageBreak
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()

    # 自定义样式
    title_style = ParagraphStyle(
        'CustomTitle', parent=styles['Title'],
        fontSize=22, spaceAfter=6, alignment=TA_CENTER,
        textColor=HexColor('#1E40AF')
    )
    date_style = ParagraphStyle(
        'DateStyle', parent=styles['Normal'],
        fontSize=9, textColor=HexColor('#888888'),
        alignment=TA_CENTER, spaceAfter=20
    )
    q_style = ParagraphStyle(
        'QStyle', parent=styles['Heading2'],
        fontSize=13, spaceBefore=12, spaceAfter=4,
        textColor=HexColor('#1E40AF'),
        borderWidth=0, borderPadding=0
    )
    a_style = ParagraphStyle(
        'AStyle', parent=styles['Normal'],
        fontSize=11, leading=16, spaceAfter=8,
        textColor=HexColor('#333333')
    )
    source_style = ParagraphStyle(
        'SourceStyle', parent=styles['Normal'],
        fontSize=9, textColor=HexColor('#888888'),
        leftIndent=12, spaceAfter=12
    )
    footer_style = ParagraphStyle(
        'FooterStyle', parent=styles['Normal'],
        fontSize=9, textColor=HexColor('#AAAAAA'),
        alignment=TA_CENTER, spaceBefore=20
    )

    # 构建内容
    elements = []

    # 标题
    elements.append(Paragraph(title, title_style))

    # 日期
    date_str = f"导出日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}"
    elements.append(Paragraph(date_str, date_style))

    # 分隔线
    elements.append(HRFlowable(width="100%", thickness=1, color=HexColor('#CCCCCC')))
    elements.append(Spacer(1, 12))

    # 问答内容
    for idx, item in enumerate(items, 1):
        elements.append(Paragraph(f"Q{idx}: {item.question}", q_style))
        elements.append(Paragraph(item.answer, a_style))

        if item.sources:
            sources_text = f"来源：{'、'.join(item.sources)}"
            elements.append(Paragraph(sources_text, source_style))

        elements.append(HRFlowable(width="60%", thickness=0.5, color=HexColor('#EEEEEE')))
        elements.append(Spacer(1, 8))

    # 页脚
    elements.append(Spacer(1, 20))
    elements.append(HRFlowable(width="100%", thickness=1, color=HexColor('#CCCCCC')))
    elements.append(Spacer(1, 8))
    elements.append(Paragraph("—— 由 LocalRAG-CS 智能客服系统生成 ——", footer_style))

    doc.build(elements)
    buf.seek(0)
    return buf.getvalue()


# ========== RTF 降级导出 ==========

def _export_rtf(items: List[ExportItem], title: str) -> bytes:
    """
    导出为 RTF 格式（零依赖降级方案）

    Args:
        items: 问答项列表
        title: 报告标题

    Returns:
        RTF 文档字节流
    """
    lines = []
    lines.append("{\\rtf1\\ansi\\deff0")
    lines.append("{\\fonttbl {\\f0 SimSun;}}")
    lines.append("\\f0\\fs24")

    # 标题
    lines.append(f"\\pard\\qc\\b\\fs36 {_escape_rtf(title)}\\b0\\par")
    lines.append(f"\\pard\\qc\\fs18 导出日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}\\par")
    lines.append("\\par")

    # 分隔线
    lines.append("\\pard\\qc\\fs16 " + "─" * 50 + "\\par")
    lines.append("\\par")

    # 问答内容
    for idx, item in enumerate(items, 1):
        lines.append(f"\\pard\\b\\fs24 Q{idx}: {_escape_rtf(item.question)}\\b0\\par")
        lines.append(f"\\pard\\fs20 {_escape_rtf(item.answer)}\\par")

        if item.sources:
            sources_text = f"来源：{'、'.join(item.sources)}"
            lines.append(f"\\pard\\fi400\\fs16\\i {_escape_rtf(sources_text)}\\i0\\par")

        lines.append("\\par")
        lines.append(f"\\pard\\qc\\fs16 " + "─" * 50 + "\\par")
        lines.append("\\par")

    # 页脚
    lines.append("\\pard\\qc\\fs16 —— 由 LocalRAG-CS 智能客服系统生成 ——\\par")
    lines.append("}")

    rtf_content = "\n".join(lines)
    return rtf_content.encode('utf-8')


def _safe_filename(title: str, ext: str) -> str:
    """生成安全的文件名（仅ASCII字符）"""
    safe = "".join(c for c in title if c.isascii() and (c.isalnum() or c in " _-"))
    safe = safe.strip() or "export"
    if len(safe) > 50:
        safe = safe[:50]
    now = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f"{safe}_{now}.{ext}"


def _escape_rtf(text: str) -> str:
    """转义 RTF 特殊字符"""
    result = []
    for ch in text:
        if ord(ch) > 127:
            result.append(f"\\u{ord(ch)}?")
        elif ch == '\\':
            result.append('\\\\')
        elif ch == '{':
            result.append('\\{')
        elif ch == '}':
            result.append('\\}')
        else:
            result.append(ch)
    return ''.join(result)


# ========== 目录浏览 ==========


class BrowseDirRequest(BaseModel):
    """浏览目录请求"""
    path: str = Field(default="", description="要浏览的目录路径，空字符串返回根目录列表")


@router.post("/browse", summary="浏览服务器目录")
async def browse_directory(req: BrowseDirRequest):
    """
    浏览服务器文件系统，返回目录列表

    参数:
        path: 目录路径，空字符串返回根目录/常用目录
    """
    try:
        # 常用快捷目录
        quick_dirs = [
            {"path": "/home/zbs/localrag-cs/exports", "label": "📁 项目导出目录"},
            {"path": "/mnt/c/Users/8/Desktop", "label": "🖥️ Windows 桌面"},
            {"path": "/mnt/c/Users/8/Documents", "label": "📄 Windows 文档"},
            {"path": "/home/zbs/localrag-cs/data", "label": "📂 项目数据目录"},
        ]

        if not req.path:
            # 根视图：检测 Windows 盘符
            windows_drives = []
            mnt_path = "/mnt"
            if os.path.isdir(mnt_path):
                try:
                    for name in sorted(os.listdir(mnt_path)):
                        drive_path = os.path.join(mnt_path, name)
                        if os.path.isdir(drive_path) and len(name) == 1 and name.isalpha():
                            windows_drives.append({
                                "path": drive_path,
                                "label": f"💿 Windows {name.upper()}: 盘",
                            })
                except PermissionError:
                    pass

            # WSL 常用路径
            wsl_paths = [
                {"path": "/home", "label": "🏠 WSL /home"},
                {"path": "/etc", "label": "⚙️ WSL /etc"},
                {"path": "/opt", "label": "📦 WSL /opt"},
                {"path": "/var", "label": "📋 WSL /var"},
                {"path": "/tmp", "label": "📝 WSL /tmp"},
                {"path": "/usr/local", "label": "🔧 WSL /usr/local"},
            ]

            return {
                "current_path": "",
                "parent_path": None,
                "entries": [],
                "quick_dirs": quick_dirs,
                "windows_drives": windows_drives,
                "wsl_paths": wsl_paths,
            }

        path = os.path.abspath(req.path)
        if not os.path.exists(path):
            raise HTTPException(status_code=404, detail=f"目录不存在: {path}")
        if not os.path.isdir(path):
            raise HTTPException(status_code=400, detail=f"不是目录: {path}")

        entries = []
        try:
            for name in sorted(os.listdir(path)):
                full = os.path.join(path, name)
                try:
                    is_dir = os.path.isdir(full)
                    stat = os.stat(full)
                    entries.append({
                        "name": name,
                        "path": full,
                        "is_dir": is_dir,
                        "size": stat.st_size if not is_dir else 0,
                        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    })
                except PermissionError:
                    entries.append({
                        "name": name,
                        "path": full,
                        "is_dir": False,
                        "size": 0,
                        "modified": "",
                    })
        except PermissionError:
            raise HTTPException(status_code=403, detail=f"无权限访问: {path}")

        parent_path = os.path.dirname(path) if path != "/" else None

        return {
            "current_path": path,
            "parent_path": parent_path,
            "entries": entries,
            "quick_dirs": quick_dirs,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"浏览目录失败: {e}")
        raise HTTPException(status_code=500, detail=f"浏览目录失败: {str(e)}")


# ========== API 路由 ==========


@router.post("/pdf", summary="导出问答报告为 PDF")
async def export_pdf(req: ExportRequest):
    """
    导出会话问答记录为 PDF 文件

    参数:
        session_id: 会话ID
        title: 报告标题（可选，默认"问答报告"）
        include_sources: 是否包含来源引用（可选，默认true）
    """
    try:
        # 检查会话是否存在
        session_data = session_manager.get_session(req.session_id)
        if not session_data:
            raise HTTPException(status_code=404, detail="会话不存在")

        # 准备数据
        items = _prepare_export_data(req.session_id, req.include_sources)
        if not items:
            raise HTTPException(status_code=400, detail="会话中没有问答记录")

        # 生成 PDF
        try:
            pdf_bytes = _export_pdf(items, req.title)
            content_type = "application/pdf"
            filename = f"{req.title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        except Exception as e:
            logger.warning(f"PDF 生成失败，降级为 RTF: {e}")
            pdf_bytes = _export_rtf(items, req.title)
            content_type = "application/rtf"
            filename = f"{req.title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.rtf"

        safe_filename = _safe_filename(req.title, "pdf")

        # 如果指定了保存路径，写入文件
        if req.save_path:
            save_dir = os.path.abspath(req.save_path)
            os.makedirs(save_dir, exist_ok=True)
            filepath = os.path.join(save_dir, safe_filename)
            with open(filepath, "wb") as f:
                f.write(pdf_bytes)
            logger.info(f"PDF 已保存到: {filepath}")
            return {"message": "导出成功", "filepath": filepath, "filename": safe_filename}

        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type=content_type,
            headers={
                "Content-Disposition": f'attachment; filename="{safe_filename}"',
                "Content-Length": str(len(pdf_bytes))
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"导出 PDF 失败: {e}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@router.post("/docx", summary="导出问答报告为 Word")
async def export_docx(req: ExportRequest):
    """
    导出会话问答记录为 Word 文件

    参数:
        session_id: 会话ID
        title: 报告标题（可选，默认"问答报告"）
        include_sources: 是否包含来源引用（可选，默认true）
    """
    try:
        # 检查会话是否存在
        session_data = session_manager.get_session(req.session_id)
        if not session_data:
            raise HTTPException(status_code=404, detail="会话不存在")

        # 准备数据
        items = _prepare_export_data(req.session_id, req.include_sources)
        if not items:
            raise HTTPException(status_code=400, detail="会话中没有问答记录")

        # 生成 Word
        try:
            docx_bytes = _export_docx(items, req.title)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        except Exception as e:
            logger.warning(f"Word 生成失败，降级为 RTF: {e}")
            docx_bytes = _export_rtf(items, req.title)
            content_type = "application/rtf"

        safe_filename = _safe_filename(req.title, "docx")

        # 如果指定了保存路径，写入文件
        if req.save_path:
            save_dir = os.path.abspath(req.save_path)
            os.makedirs(save_dir, exist_ok=True)
            filepath = os.path.join(save_dir, safe_filename)
            with open(filepath, "wb") as f:
                f.write(docx_bytes)
            logger.info(f"Word 已保存到: {filepath}")
            return {"message": "导出成功", "filepath": filepath, "filename": safe_filename}

        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type=content_type,
            headers={
                "Content-Disposition": f'attachment; filename="{safe_filename}"',
                "Content-Length": str(len(docx_bytes))
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"导出 Word 失败: {e}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")
